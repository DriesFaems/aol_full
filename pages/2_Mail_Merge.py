import io
import zipfile
from typing import List, Dict, Tuple

import streamlit as st
try:
    from docx import Document
    HAS_DOCX = True
except Exception:
    Document = None
    HAS_DOCX = False
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pandas as pd
from openpyxl import load_workbook
from utils import set_branding

# Apply Branding
set_branding()

def _detect_extension(filename: str) -> str:
    return filename.split(".")[-1].lower()

def parse_tabular_upload(uploaded) -> List[Dict]:
    """Parse CSV or Excel uploads into structured student records.

    Expected columns:
    - First name, Last name, Email address, Final Assignment Grade
    - Any number of dimension columns named like "DimensionName - Score", "DimensionName - Notes"
    """
    ext = _detect_extension(uploaded.name)
    try:
        if ext in ("xls", "xlsx"):
            xfile = pd.ExcelFile(uploaded, engine="openpyxl")
            if "Evaluation Form" in xfile.sheet_names:
                df = pd.read_excel(xfile, sheet_name="Evaluation Form", engine="openpyxl")
            else:
                df = pd.read_excel(xfile, sheet_name=0, engine="openpyxl")
        else:
            df = pd.read_csv(uploaded)
    except Exception as e:
        raise ValueError(f"Could not read tabular file: {e}")

    required = ["First name", "Last name", "Email address", "Final Assignment Grade"]
    missing = [c for c in required if c not in df.columns]
    if missing:
        raise ValueError(f"Missing required columns: {', '.join(missing)}")

    # detect dimension groups
    dimension_names = set()
    for col in df.columns:
        if " - " in col:
            name, _ = col.split(" - ", 1)
            dimension_names.add(name.strip())

    # Try to read Rubric sheet and Assignment Info sheet from the same workbook to enrich rubric text and weights
    rubric_map: Dict[str, str] = {}
    rubric_levels_map: Dict[str, Dict[str, str]] = {}
    rubric_weights_map: Dict[str, str] = {}
    assignment_weights_map: Dict[str, str] = {}
    if ext in ("xls", "xlsx"):
        try:
            xfile = pd.ExcelFile(uploaded, engine="openpyxl")
            if "Rubric" in xfile.sheet_names:
                rubric_df = pd.read_excel(xfile, sheet_name="Rubric", engine="openpyxl")
                if "Dimension" in rubric_df.columns:
                    desc_col = None
                    for candidate in ("Rubric", "Description", "Notes"):
                        if candidate in rubric_df.columns:
                            desc_col = candidate
                            break
                    level_names = ["Very Good", "Good", "Satisfactory", "Unsatisfactory"]
                    col_map = {}
                    for col in rubric_df.columns:
                        col_norm = col.strip().lower().replace(" ", "")
                        for lvl in level_names:
                            lvl_norm = lvl.strip().lower().replace(" ", "")
                            if lvl == "Unsatisfactory":
                                if (
                                    "satisfactory" in col_norm
                                    and not any(x in col_norm for x in ["verysatisfactory", "goodsatisfactory", "verygood", "good"])
                                ):
                                    col_map[lvl] = col
                            elif col_norm == lvl_norm:
                                col_map[lvl] = col
                    weight_col = None
                    for candidate in ("Weight (%)", "Weight", "Dimension Weight"):
                        if candidate in rubric_df.columns:
                            weight_col = candidate
                            break
                    for _, rrow in rubric_df.iterrows():
                        dim = str(rrow.get("Dimension", "")).strip()
                        if not dim:
                            continue
                        rubric_text = rrow.get(desc_col, "") or ""
                        rubric_map[dim] = str(rubric_text)
                        rubric_levels_map[dim] = {lvl: str(rrow.get(col_map.get(lvl, ""), "")) for lvl in level_names}
                        rubric_weights_map[dim] = str(rrow.get(weight_col, "")) if weight_col else ""
            if "Assignment Info" in xfile.sheet_names:
                try:
                    assign_df = pd.read_excel(xfile, sheet_name="Assignment Info", engine="openpyxl")
                    # If the sheet already has a 'Dimension' column, read per-dimension weights
                    if "Dimension" in assign_df.columns:
                        weight_candidates = [c for c in assign_df.columns if c.strip().lower() in ("weight(%)", "weight", "weight (%)", "dimension weight")]
                        weight_col_assign = weight_candidates[0] if weight_candidates else None
                        for _, arow in assign_df.iterrows():
                            dim_name = str(arow.get("Dimension", "")).strip()
                            if not dim_name:
                                continue
                            if weight_col_assign:
                                assignment_weights_map[dim_name] = str(arow.get(weight_col_assign, ""))
                    else:
                        # If the sheet is key/value pairs, pivot into a dict and keep as general assignment info.
                        # This is not per-dimension; we inform the user in the UI to provide a column format for per-dimension weights.
                        # Build dict of key->value for potential use (e.g., assignment weight)
                        if assign_df.shape[1] >= 2:
                            key_col = assign_df.columns[0]
                            val_col = assign_df.columns[1]
                            assign_kv = {}
                            for _, arow in assign_df.iterrows():
                                k = str(arow.get(key_col, "")).strip()
                                v = arow.get(val_col, "")
                                if k:
                                    assign_kv[k] = v
                            # If any keys match a dimension name, use them
                            for k, v in assign_kv.items():
                                if k in rubric_weights_map:
                                    assignment_weights_map[k] = str(v)
                except Exception:
                    assignment_weights_map = {}
        except Exception:
            # ignore workbook read errors
            pass

    students: List[Dict] = []
    for _, row in df.iterrows():
        student = {
            "first_name": row.get("First name", ""),
            "last_name": row.get("Last name", ""),
            "email": row.get("Email address", ""),
            "final_grade": row.get("Final Assignment Grade", ""),
            "dimensions": [],
        }
        for dim in sorted(dimension_names):
            score_col = f"{dim} - Score"
            notes_col = f"{dim} - Notes"
            # determine weight: prefer Assignment Info, then Rubric sheet, else empty
            weight = assignment_weights_map.get(dim, rubric_weights_map.get(dim, ""))
            student["dimensions"].append(
                {
                    "name": dim,
                    "grade": row.get(score_col, ""),
                    "notes": row.get(notes_col, ""),
                    "rubric_levels": rubric_levels_map.get(dim, {}),
                    "weight": weight,
                }
            )
        students.append(student)
    return students


def parse_docx_upload(uploaded) -> List[Dict]:
    """A very permissive .docx parser. It tries to split the document into blocks per student.

    This will work best when the uploaded document uses clear separators like blank lines or headings
    such as 'Student: Name' or 'Name:'. If parsing fails, prefer CSV/Excel upload.
    """
    try:
        doc = Document(io.BytesIO(uploaded.getvalue()))
    except Exception as e:
        raise ValueError(f"Could not read .docx file: {e}")

    text = "\n\n".join(p.text for p in doc.paragraphs if p.text and p.text.strip())
    # split into blocks by double newlines
    blocks = [b.strip() for b in text.split("\n\n") if b.strip()]
    students: List[Dict] = []
    for block in blocks:
        lines = [l.strip() for l in block.splitlines() if l.strip()]
        # attempt to find name/email
        name = None
        email = None
        final_grade = ""
        dimensions: List[Dict] = []
        current_dim = None
        for line in lines:
            low = line.lower()
            if low.startswith("student:") or low.startswith("name:"):
                name = line.split(":", 1)[1].strip()
                continue
            if "@" in line and "." in line:
                email = line.strip()
                continue
            if low.startswith("final") and "grade" in low:
                # e.g. Final grade: 7.5
                parts = line.split(":", 1)
                if len(parts) > 1:
                    final_grade = parts[1].strip()
                continue
            # dimension lines heuristics: "DimensionName - Grade: X" or "DimensionName: Grade X"
            if " - " in line and ("score" in low or "grade" in low or ":" in line):
                # split into dim and remainder
                name_part, rest = line.split(" - ", 1)
                current_dim = {"name": name_part.strip(), "grade": rest.strip(), "notes": "", "rubric": ""}
                dimensions.append(current_dim)
                continue
            if ":" in line and not current_dim:
                # maybe "DimensionName: Grade X"
                left, right = line.split(":", 1)
                if any(k in left.lower() for k in ("dimension", "criterion")) or any(k in right.lower() for k in ("grade", "score")):
                    current_dim = {"name": left.strip(), "grade": right.strip(), "notes": "", "rubric": ""}
                    dimensions.append(current_dim)
                    continue
            # lines that look like notes or rubric
            if current_dim is not None:
                if low.startswith("notes") or low.startswith("justification"):
                    parts = line.split(":", 1)
                    current_dim["notes"] = parts[1].strip() if len(parts) > 1 else ""
                    continue
                if low.startswith("rubric"):
                    parts = line.split(":", 1)
                    current_dim["rubric"] = parts[1].strip() if len(parts) > 1 else ""
                    continue

        # If we parsed a 'name' that might include first and last name
        first_name = ""
        last_name = ""
        if name:
            parts = name.split()
            first_name = parts[0]
            last_name = " ".join(parts[1:]) if len(parts) > 1 else ""

        students.append(
            {
                "first_name": first_name,
                "last_name": last_name,
                "email": email or "",
                "final_grade": final_grade,
                "dimensions": dimensions,
            }
        )
    return students


def generate_combined_doc(students: List[Dict]) -> bytes:
    doc = Document()
    doc.add_heading("Assignment Evaluations", level=1)
    for student in students:
        name = f"{student.get('first_name','').strip()} {student.get('last_name','').strip()}".strip()
        doc.add_heading(name or student.get("email", "Unnamed Student"), level=2)
        if student.get("email"):
            doc.add_paragraph(f"Email: {student.get('email')}")
        for dim in student.get("dimensions", []):
            doc.add_heading(dim.get("name", "Dimension"), level=3)
            table = doc.add_table(rows=2, cols=2)
            table.style = "Table Grid"
            table.cell(0, 0).text = "Grade"
            table.cell(0, 1).text = str(dim.get("grade", ""))
            table.cell(1, 0).text = "Notes"
            table.cell(1, 1).text = str(dim.get("notes", ""))
            # Add weight
            if dim.get("weight"):
                doc.add_paragraph(f"Dimension weight: {dim['weight']}")
            # Add rubric levels if available (all four levels)
            rubric_levels = dim.get("rubric_levels", {})
            if rubric_levels:
                doc.add_paragraph("Rubric levels:", style="Heading 4")
                for level in ["Very Good", "Good", "Satisfactory", "Unsatisfactory"]:
                    doc.add_paragraph(f"{level}: {rubric_levels.get(level, '')}")
        doc.add_paragraph(f"Final grade: {student.get('final_grade', '')}")
        doc.add_page_break()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_individual_zip(students: List[Dict]) -> bytes:
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        for student in students:
            doc = Document()
            name = f"{student.get('first_name','').strip()} {student.get('last_name','').strip()}".strip()
            doc.add_heading(name or student.get('email', 'Unnamed Student'), level=1)
            if student.get('email'):
                doc.add_paragraph(f"Email: {student.get('email')}")
            for dim in student.get('dimensions', []):
                doc.add_heading(dim.get('name', 'Dimension'), level=2)
                doc.add_paragraph(f"Grade: {dim.get('grade', '')}")
                doc.add_paragraph(f"Notes: {dim.get('notes', '')}")
                # Add weight
                if dim.get("weight"):
                    doc.add_paragraph(f"Dimension weight: {dim['weight']}")
                rubric_levels = dim.get("rubric_levels", {})
                if rubric_levels:
                    doc.add_paragraph("Rubric levels:", style="Heading 4")
                    for level in ["Very Good", "Good", "Satisfactory", "Unsatisfactory"]:
                        doc.add_paragraph(f"{level}: {rubric_levels.get(level, '')}")
            doc.add_paragraph(f"Final grade: {student.get('final_grade', '')}")

            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            filename = (name or student.get('email', 'student')).replace(" ", "_") + ".docx"
            zf.writestr(filename, buf.getvalue())
    zip_buffer.seek(0)
    return zip_buffer.getvalue()


def sample_csv_template() -> bytes:
    df = pd.DataFrame(
        [
            {
                "First name": "Alice",
                "Last name": "Example",
                "Email address": "alice@example.com",
                "Final Assignment Grade": 8.5,
                "Understanding - Score": 8,
                "Understanding - Notes": "Clear reasoning",
                "Understanding - Rubric": "Meets expectations",
            }
        ]
    )
    buf = io.BytesIO()
    df.to_csv(buf, index=False)
    buf.seek(0)
    return buf.getvalue()


def _safe_field_name(name: str) -> str:
    # create a compact merge field name without spaces
    return (
        name.replace(" - ", "_")
        .replace(" ", "_")
        .replace("/", "_")
        .replace("\\", "_")
    )


def build_merge_dataframe(students: List[Dict]) -> pd.DataFrame:
    # collect dimension names
    dim_names = []
    if students:
        dim_names = [d["name"] for d in students[0].get("dimensions", [])]

    rows = []
    for s in students:
        row = {
            "FirstName": s.get("first_name", ""),
            "LastName": s.get("last_name", ""),
            "Email": s.get("email", ""),
            "FinalGrade": s.get("final_grade", ""),
        }
        for dim in dim_names:
            safe = _safe_field_name(dim)
            # find matching dimension in student
            found = next((dd for dd in s.get("dimensions", []) if dd.get("name") == dim), None)
            row[f"{safe}_Score"] = found.get("grade", "") if found else ""
            row[f"{safe}_Notes"] = found.get("notes", "") if found else ""
            row[f"{safe}_Weight"] = found.get("weight", "") if found else ""
        rows.append(row)
    df = pd.DataFrame(rows)
    return df


def _add_mergefield(paragraph, field_name: str):
    # Insert a simple MERGEFIELD into a paragraph using fldSimple
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), f'MERGEFIELD "{field_name}"')
    paragraph._p.append(fld)
    return paragraph


def create_mailmerge_docx(sample_df: pd.DataFrame) -> bytes:
    if not HAS_DOCX:
        raise RuntimeError("python-docx is not available")

    doc = Document()
    doc.add_heading("Assignment Evaluation - Mail Merge Template", level=1)
    p = doc.add_paragraph()
    p.add_run("Dear ")
    _add_mergefield(p, "FirstName")
    p.add_run(" ")
    _add_mergefield(p, "LastName")
    p.add_run(",")

    doc.add_paragraph("")
    p2 = doc.add_paragraph()
    p2.add_run("Final grade: ")
    _add_mergefield(p2, "FinalGrade")

    # For each dimension, include fields for Score, Notes, Rubric, and rubric levels
    # Derive dimensions from columns in sample_df
    dim_prefixes = []
    rubric_levels_map = getattr(sample_df, "rubric_levels_map", {})
    for col in sample_df.columns:
        if col.endswith("_Score"):
            prefix = col[: -len("_Score")]
            dim_prefixes.append(prefix)

    for prefix in dim_prefixes:
        human = prefix.replace("_", " ")
        doc.add_paragraph(f"{human}", style="Heading 2")
        p_score = doc.add_paragraph()
        p_score.add_run("Grade: ")
        _add_mergefield(p_score, f"{prefix}_Score")
        p_notes = doc.add_paragraph()
        p_notes.add_run("Notes: ")
        _add_mergefield(p_notes, f"{prefix}_Notes")
        # Add weight
        doc.add_paragraph("Dimension weight: ", style="Heading 4")
        _add_mergefield(doc.add_paragraph(), f"{prefix}_Weight")
        # Add rubric levels if available (all four levels)
        if rubric_levels_map and prefix in rubric_levels_map:
            doc.add_paragraph("Rubric levels:", style="Heading 4")
            for level in ["Very Good", "Good", "Satisfactory", "Unsatisfactory"]:
                doc.add_paragraph(f"{level}: {rubric_levels_map[prefix].get(level, '')}")


    # Sign-off line
    doc.add_paragraph("")
    doc.add_paragraph("Kind regards,")
    doc.add_paragraph("[Your name]")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def generate_mailmerge_zip(students: List[Dict]) -> bytes:
    df = build_merge_dataframe(students)
    # Attach rubric_levels_map to df for use in template generation
    rubric_levels_map = {}
    if students and students[0].get("dimensions"):
        for dim in students[0]["dimensions"]:
            safe = _safe_field_name(dim["name"])
            rubric_levels_map[safe] = dim.get("rubric_levels", {})
    df.rubric_levels_map = rubric_levels_map
    csv_buf = io.BytesIO()
    df.to_csv(csv_buf, index=False)
    csv_buf.seek(0)

    docx_bytes = create_mailmerge_docx(df)

    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("mailmerge_data.csv", csv_buf.getvalue())
        zf.writestr("mailmerge_template.docx", docx_bytes)
    zip_buf.seek(0)
    return zip_buf.getvalue()


def main():
    st.title("Assignment Evaluations → Mail-merge Word Generator")
    st.write(
        "Upload a CSV/Excel export of evaluations (recommended) or a Word document with evaluations. "
        "The app will create Word documents containing each student's evaluation (dimensions, notes, rubric, final grade)."
    )

    st.markdown("**Step 1 — Download a sample CSV template (recommended)**")
    st.download_button("Download sample CSV", data=sample_csv_template(), file_name="sample_evaluations.csv")

    uploaded = st.file_uploader("Upload evaluation file", type=["csv", "xlsx", "xls", "docx"])
    if not uploaded:
        st.info("Upload a file to generate personalized evaluation documents.")
        return

    try:
        ext = _detect_extension(uploaded.name)
        if ext in ("csv", "xls", "xlsx"):
            students = parse_tabular_upload(uploaded)
        elif ext == "docx":
            students = parse_docx_upload(uploaded)
        else:
            st.error("Unsupported file type")
            return
    except Exception as e:
        st.error(f"Could not parse upload: {e}")
        return

    if not students:
        st.warning("No student evaluations were detected in the uploaded file.")
        return

    st.success(f"Parsed {len(students)} student evaluation(s).")

    if st.checkbox("Show parsed data (for debugging)"):
        st.json(students)

    st.markdown("**Output options**")
    combined = st.button("Generate combined Word document (.docx)")
    individual = st.button("Generate individual .docx files and download ZIP")
    mailmerge = st.button("Generate mail-merge template + data (ZIP)")

    if combined:
        docx_bytes = generate_combined_doc(students)
        st.download_button("Download combined .docx", data=docx_bytes, file_name="evaluations_combined.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    if individual:
        zip_bytes = generate_individual_zip(students)
        st.download_button("Download ZIP with individual docs", data=zip_bytes, file_name="evaluations_individual.zip", mime="application/zip")

    if mailmerge:
        # If python-docx is not installed, still offer CSV and explain how to install python-docx
        if not HAS_DOCX:
            st.error("`python-docx` is not installed in this environment, so a Word template cannot be created.")
            # provide CSV data for mail merge
            df = build_merge_dataframe(students)
            csv_buf = io.BytesIO()
            df.to_csv(csv_buf, index=False)
            csv_buf.seek(0)
            st.download_button("Download CSV for mail-merge", data=csv_buf.getvalue(), file_name="mailmerge_data.csv", mime="text/csv")
            st.info("Install `python-docx` (python -m pip install python-docx) and regenerate to get the Word template included in a ZIP.")
        else:
            try:
                zip_bytes = generate_mailmerge_zip(students)
            except Exception as e:
                st.error(f"Failed to create mail-merge package: {e}")
            else:
                st.download_button("Download mail-merge ZIP (template + CSV)", data=zip_bytes, file_name="mailmerge_package.zip", mime="application/zip")


if __name__ == "__main__":
    main()
